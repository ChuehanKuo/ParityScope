"""Update the ParityScope Business Plan with NHANES investigation findings."""

import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

INPUT = "/home/user/ParityScopeWebsite/ParityScope_Business_Plan (8).docx"
OUTPUT = "/home/user/ParityScopeWebsite/ParityScope_Business_Plan (9).docx"

doc = docx.Document(INPUT)


def find_paragraph_index(doc, text_fragment):
    """Find the index of the paragraph containing the given text."""
    for i, p in enumerate(doc.paragraphs):
        if text_fragment in p.text:
            return i
    return None


def add_heading(doc, text, level, index):
    """Insert a heading at a specific paragraph index."""
    new_para = doc.paragraphs[index].insert_paragraph_before(text)
    # Copy heading style from an existing heading of the same level
    for p in doc.paragraphs:
        if p.style and p.style.name == f"Heading {level}":
            new_para.style = p.style
            break
    return new_para


def add_paragraph(doc, text, index, bold_prefix=None):
    """Insert a paragraph at a specific paragraph index."""
    new_para = doc.paragraphs[index].insert_paragraph_before("")
    if bold_prefix:
        run_bold = new_para.add_run(bold_prefix)
        run_bold.bold = True
        run_normal = new_para.add_run(text)
    else:
        run = new_para.add_run(text)
    return new_para


# ---------------------------------------------------------------
# 1. Add NHANES Case Study after "OUR SOLUTION" section
#    Insert before "How It Works" heading (index 35)
# ---------------------------------------------------------------

# Find the "How It Works" heading
how_it_works_idx = find_paragraph_index(doc, "How It Works")
print(f"Found 'How It Works' at index {how_it_works_idx}")

# We'll insert BEFORE "How It Works" — a new section showing ParityScope in action
# Insert in reverse order (each insert goes before the same index, pushing others down)

insert_texts = [
    ("heading2", "ParityScope in Action: NHANES Hypertension Model"),
    ("para", "To validate ParityScope's capabilities, we audited a hypertension prediction model trained on real CDC population health data (NHANES 2017–2018, 5,216 adults). The model — a standard logistic regression with clinical features — represents the type of clinical AI commonly deployed in healthcare settings. The results demonstrate exactly the kind of hidden failures ParityScope is designed to catch."),
    ("bold", "The headline metric hides the truth. ", "The model reports 66% overall accuracy — a number that would satisfy most internal reviews. But ParityScope's per-group analysis revealed the model is fundamentally broken: it predicts zero hypertension cases for anyone under 40, and labels 100% of patients over 80 as hypertensive — regardless of their actual health status."),
    ("bold", "The root cause: a model masquerading as a classifier. ", "ParityScope's coefficient analysis showed that age accounts for 56.7% of the model's total decision weight — more than all other clinical features combined. BMI, cholesterol, smoking, diabetes, and kidney disease are effectively ignored. The model is not detecting hypertension; it is detecting age. For a patient with median features, the model flips from 'healthy' to 'hypertensive' at exactly age 58, regardless of clinical indicators."),
    ("bold", "Real patients harmed in both directions. ", "In the 18–39 age group (491 test patients, 6.5% actual prevalence), the model missed every single hypertensive patient — 32 people with a treatable condition who received no flag. In the 80+ group (111 test patients, 53.9% actual prevalence), 46% of patients were healthy but all were labeled hypertensive — driving unnecessary follow-up and resource waste."),
    ("bold", "Racial disparities compound the problem. ", "Beyond age, ParityScope detected systematic racial bias: the model detects hypertension at 77.6% for Non-Hispanic White patients but only 56.3% for Other/Multi-Racial patients — a 21-percentage-point gap in detection rates. All 8 fairness metrics across race were rated UNFAIR under EU AI Act thresholds."),
    ("bold", "The core value proposition demonstrated. ", "A standard model evaluation would report '66% accuracy, AUC 0.78' and the model would pass review. ParityScope reveals it is a crude age threshold wearing a lab coat — catching model failures that overall metrics hide. This is not a theoretical risk; it is a validated finding on real U.S. population health data."),
]

# Insert in reverse so indices don't shift
for item in reversed(insert_texts):
    if item[0] == "heading2":
        add_heading(doc, item[1], 2, how_it_works_idx)
    elif item[0] == "para":
        add_paragraph(doc, item[1], how_it_works_idx)
    elif item[0] == "bold":
        add_paragraph(doc, item[2], how_it_works_idx, bold_prefix=item[1])

print("Added NHANES case study section")

# ---------------------------------------------------------------
# 2. Add the "Hidden Failures" pitch to the Value Proposition
#    in the Key Differentiators section
# ---------------------------------------------------------------

# Find the last differentiator paragraph (Consulting-led)
consulting_idx = find_paragraph_index(doc, "Consulting-led product model")
print(f"Found 'Consulting-led' at index {consulting_idx}")

# Insert a new differentiator after it
insert_idx = consulting_idx + 1
add_paragraph(
    doc,
    " Standard model evaluation metrics — overall accuracy, AUC, F1 — average performance across all patients. "
    "A model can report 66% accuracy while failing catastrophically for entire demographic groups. "
    "ParityScope breaks open the single number and reveals who the model works for and who it fails. "
    "In healthcare, the patients it fails can die.",
    insert_idx,
    bold_prefix="Catches failures that overall metrics hide."
)

print("Added 'hidden failures' differentiator")

# ---------------------------------------------------------------
# 3. Add NHANES finding to the Problem section's bias cases table
#    Find the paragraph after the table (The Governance Gap heading)
# ---------------------------------------------------------------

governance_gap_idx = find_paragraph_index(doc, "The Governance Gap")
print(f"Found 'The Governance Gap' at index {governance_gap_idx}")

# Insert before "The Governance Gap"
add_paragraph(
    doc,
    "ParityScope's own audit of a hypertension prediction model trained on CDC NHANES data (2017–2018) found the model "
    "reports 66% overall accuracy but predicts zero cases for patients under 40 and 100% positive for patients over 80. "
    "Age accounts for 56.7% of the model's decisions — all other clinical features are effectively ignored. "
    "Racial disparities showed a 21-percentage-point gap in detection rates. The model passed standard evaluation metrics "
    "but failed 19 of 24 fairness metrics under EU AI Act thresholds.",
    governance_gap_idx,
    bold_prefix="NHANES Hypertension Model (ParityScope Audit, 2026): "
)

print("Added NHANES to bias cases")

# ---------------------------------------------------------------
# Save
# ---------------------------------------------------------------
doc.save(OUTPUT)
print(f"\nSaved updated business plan to: {OUTPUT}")
